import re
import shutil
from pathlib import Path
from typing import Dict, List


def sanitize_filename(name: str) -> str:
    sanitized = re.sub(r'[<>:"/\\|?*]', '_', name)
    sanitized = sanitized.strip(' .')
    sanitized = re.sub(r'_+', '_', sanitized)
    return sanitized


def duplicate_document(source_path: Path, company_name: str) -> Path:
    if not source_path.exists():
        raise FileNotFoundError(f"Source file not found: {source_path}")
    
    sanitized_company = sanitize_filename(company_name)
    if not sanitized_company or sanitized_company == "unknown":
        sanitized_company = "customized"
    
    stem = source_path.stem
    suffix = source_path.suffix
    new_name = f"{stem}_{sanitized_company}{suffix}"
    new_path = source_path.parent / new_name
    
    counter = 1
    original_new_path = new_path
    while new_path.exists():
        new_name = f"{stem}_{sanitized_company}_{counter}{suffix}"
        new_path = source_path.parent / new_name
        counter += 1
        if counter > 1000:  # Safety limit
            raise RuntimeError(f"Could not generate unique filename for {original_new_path}")
    
    shutil.copy2(source_path, new_path)
    return new_path


def apply_replacements_to_docx(file_path: Path, replacements: Dict[str, List[Dict[str, str]]]) -> None:
    from docx import Document
    
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        doc = Document(str(file_path))
        
        all_replacements = []
        all_replacements.extend(replacements.get("role_replacements", []))
        all_replacements.extend(replacements.get("skill_replacements", []))
        
        if not all_replacements:
            return
        
        def replace_in_paragraph(paragraph, old_text: str, new_text: str) -> bool:
            if old_text not in paragraph.text:
                return False
            
            full_text = paragraph.text
            if old_text in full_text:
                paragraph.clear()
                parts = full_text.split(old_text, 1)
                if len(parts) == 2:
                    if parts[0]:
                        paragraph.add_run(parts[0])
                    paragraph.add_run(new_text)
                    if parts[1]:
                        paragraph.add_run(parts[1])
                    return True
            return False
        
        for replacement in all_replacements:
            old_text = replacement.get("from", "")
            new_text = replacement.get("to", "")
            
            if not old_text or old_text == new_text:
                continue
            
            for paragraph in doc.paragraphs:
                replace_in_paragraph(paragraph, old_text, new_text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph, old_text, new_text)
            
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        replace_in_paragraph(paragraph, old_text, new_text)
                    for table in section.header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    replace_in_paragraph(paragraph, old_text, new_text)
                
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        replace_in_paragraph(paragraph, old_text, new_text)
                    for table in section.footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    replace_in_paragraph(paragraph, old_text, new_text)
        
        doc.save(str(file_path))
        
    except Exception as e:
        raise RuntimeError(f"Failed to edit DOCX file {file_path}: {e}") from e


def apply_replacements_to_odt(file_path: Path, replacements: Dict[str, List[Dict[str, str]]]) -> None:
    from odf.opendocument import load
    from odf.text import P, Span
    
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        doc = load(str(file_path))
        
        all_replacements = []
        all_replacements.extend(replacements.get("role_replacements", []))
        all_replacements.extend(replacements.get("skill_replacements", []))
        
        if not all_replacements:
            return
        
        def replace_in_text_element(element, old_text: str, new_text: str) -> bool:
            text_nodes = []
            for node in element.childNodes:
                if node.nodeType == node.TEXT_NODE:
                    text_nodes.append(node)
            
            full_text = "".join([node.data for node in text_nodes])
            
            if old_text in full_text:
                for node in text_nodes:
                    element.removeChild(node)
                
                parts = full_text.split(old_text, 1)
                if len(parts) == 2:
                    if parts[0]:
                        text_node = doc.createTextNode(parts[0])
                        element.appendChild(text_node)
                    new_text_node = doc.createTextNode(new_text)
                    element.appendChild(new_text_node)
                    if parts[1]:
                        remaining_node = doc.createTextNode(parts[1])
                        element.appendChild(remaining_node)
                    return True
            return False
        
        for replacement in all_replacements:
            old_text = replacement.get("from", "")
            new_text = replacement.get("to", "")
            
            if not old_text or old_text == new_text:
                continue
            
            for paragraph in doc.getElementsByType(P):
                replace_in_text_element(paragraph, old_text, new_text)
        
        doc.save(str(file_path))
        
    except Exception as e:
        raise RuntimeError(f"Failed to edit ODT file {file_path}: {e}") from e


def apply_replacements(file_path: Path, replacements: Dict[str, List[Dict[str, str]]]) -> None:
    suffix = file_path.suffix.lower()
    
    if suffix == ".docx":
        apply_replacements_to_docx(file_path, replacements)
    elif suffix == ".odt":
        apply_replacements_to_odt(file_path, replacements)
    else:
        raise ValueError(f"Unsupported file format for editing: {suffix}. Use .docx or .odt")

