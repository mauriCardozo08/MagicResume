import os
from typing import Union

def read_docx(file_path: str) -> str:
    from docx import Document

    document = Document(file_path)
    content = []
    
    for p in document.paragraphs:
        if p.text.strip():
            content.append(p.text)

    for table in document.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                content.append(" | ".join(row_text))

    return "\n".join(content)


def read_odt(file_path: str) -> str:
    from odf.opendocument import load
    from odf.text import P
    from odf.table import Table, TableRow, TableCell

    doc = load(file_path)
    content = []

    for element in doc.getElementsByType(P):
        text_parts = []
        for node in element.childNodes:
            if node.nodeType == node.TEXT_NODE:
                text_parts.append(node.data)
        paragraph_text = "".join(text_parts).strip()
        if paragraph_text:
            content.append(paragraph_text)

    for table in doc.getElementsByType(Table):
        for row in table.getElementsByType(TableRow):
            row_text = []
            for cell in row.getElementsByType(TableCell):
                cell_text_parts = []
                for p_node in cell.getElementsByType(P):
                    for node in p_node.childNodes:
                        if node.nodeType == node.TEXT_NODE:
                           cell_text_parts.append(node.data)
                
                cell_text = "".join(cell_text_parts).strip()
                if cell_text:
                    row_text.append(cell_text)
            
            if row_text:
                content.append(" | ".join(row_text))

    return "\n".join(content)


def read_txt(file_path: str) -> str:
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def read_document_as_text(file_path: str) -> str:
    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".docx":
        return read_docx(file_path)
    elif extension == ".odt":
        return read_odt(file_path)
    elif extension == ".txt":
        return read_txt(file_path)
    else:
        raise ValueError("Unsupported file format. Use .docx, .odt, or .txt")
